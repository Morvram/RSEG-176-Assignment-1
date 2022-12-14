#Joshua Mandell
#10/25/2022
#Document Conversion code for first AWS assignment

import aspose.words as aw
import os

os.chdir("C:/Users/joshu/Dropbox/Data Science Graduate Program/Brandeis - Master's Strategic Analytics/RSEG 176/Assignment 1/RSEG-176-Assignment-1")

def convertDoc (document, outform): #input: 'document' is path to .pdf document. 'outform' is the file extension.
    doc = removeWatermark(aw.Document(document))
    doc.save(document.split(".")[0] + outform)
    
def test_pdfToRTF ():
    os.chdir("C:/Users/joshu/Dropbox/Data Science Graduate Program/Brandeis - Master's Strategic Analytics/RSEG 176/Assignment 1/RSEG-176-Assignment-1") #obviously change to whatever your working directory is.
    convertDoc("example_pdf_syllabus.pdf", ".rtf")
    
def test_rtfToDocx(): #updated 11/1/2022 to delete watermark.
    convertDoc("example_pdf_syllabus.rtf", ".docx")
    d = docx.Document("example_pdf_syllabus.docx")
    delete_paragraph(d.paragraphs[0])
    d.save("example_pdf_syllabus.docx")
    
#JM 10/25/2022: tested both of the 'test' methods to make sure they work at creating new files. They do work - copying all the pdf text. However, there is a watermark (in the form of red text on the first page) added by aspose.words
    #TODO use file manipulation to remove the watermark after creating the file?
#Also, note, this was tested with a pdf with searchable text. If using a pdf without searchable text, unsure if the conversion would work properly. OCR may be necessary as an intervention (perhaps more ambitious than required for our minimum viable prototype)

###10/31/2022: watermark (remove? change?)
##By default, there is a watermark which says: "Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd."

def removeWatermark (doc): #takes a aw.Document object
    options = aw.TextWatermarkOptions()
    options.font_size = 2
    #options.Color = Color.White
    #options.IsSemitransparent = True
    doc.watermark.set_text("_")
    return(doc)
#The above method does NOT remove the "evaluation only" bit. So this is not the same thing as the watermark.set_text()

##Extract text from a PDF

import PyPDF2

def extractPDFText(pdfloc): #pdfloc = filepath of a PDF
    pdfFile = open(pdfloc, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFile)
    text = ""
    c = pdfReader.numPages
    for i in range(c):
        text += (pdfReader.getPage(i).extractText())
    return(text)
    
t = extractPDFText("example_pdf_syllabus.pdf")



#11/1/2022: trying to remove text from Docx using Python
import docx
def delete_paragraph(paragraph): #deletes the aspose.words watermark from a docx file.
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    
import PyRTF
def del_par_rtf